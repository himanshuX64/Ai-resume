"""
Resume Parser Utility
Extracts text and skills from PDF and DOCX resume files
"""
import re
from typing import List, Dict, Optional
import PyPDF2
import docx
from io import BytesIO


class ResumeParser:
    """Parse resume files and extract skills"""
    
    def __init__(self):
        """Initialize resume parser"""
        # Common skill keywords to look for
        self.skill_patterns = [
            r'\b(?:python|java|javascript|typescript|c\+\+|c#|ruby|php|swift|kotlin|go|rust)\b',
            r'\b(?:react|angular|vue|node\.?js|express|django|flask|fastapi|spring|laravel)\b',
            r'\b(?:sql|mysql|postgresql|mongodb|redis|elasticsearch|cassandra|oracle)\b',
            r'\b(?:aws|azure|gcp|docker|kubernetes|jenkins|terraform|ansible)\b',
            r'\b(?:machine learning|deep learning|nlp|computer vision|tensorflow|pytorch|scikit-learn)\b',
            r'\b(?:git|github|gitlab|bitbucket|ci/cd|devops|agile|scrum)\b',
            r'\b(?:html|css|sass|less|bootstrap|tailwind|material-ui)\b',
            r'\b(?:rest api|graphql|microservices|websockets|grpc)\b',
            r'\b(?:pandas|numpy|matplotlib|seaborn|plotly|jupyter)\b',
            r'\b(?:linux|unix|bash|shell scripting|powershell)\b'
        ]
    
    def parse_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text from PDF
        """
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    def parse_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text from DOCX
        """
        try:
            docx_file = BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    def parse_resume(self, file_content: bytes, filename: str) -> str:
        """
        Parse resume file based on extension
        
        Args:
            file_content: File content as bytes
            filename: Name of the file
            
        Returns:
            Extracted text from resume
        """
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self.parse_pdf(file_content)
        elif file_extension in ['docx', 'doc']:
            return self.parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Please upload PDF or DOCX files.")
    
    def extract_skills_from_text(self, text: str) -> List[str]:
        """
        Extract skills from resume text using pattern matching
        
        Args:
            text: Resume text content
            
        Returns:
            List of extracted skills
        """
        text_lower = text.lower()
        skills = set()
        
        # Extract skills using patterns
        for pattern in self.skill_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                skill = match.group(0)
                # Clean up the skill
                skill = skill.strip().title()
                # Special cases for proper capitalization
                skill = self._normalize_skill_name(skill)
                # Validate before adding
                if self._is_valid_skill(skill):
                    skills.add(skill)
        
        # Look for skills section
        skills_section = self._extract_skills_section(text)
        if skills_section:
            skills.update(skills_section)
        
        return sorted(list(skills))
    
    def _normalize_skill_name(self, skill: str) -> str:
        """
        Normalize skill names for proper capitalization
        
        Args:
            skill: Skill name to normalize
            
        Returns:
            Normalized skill name
        """
        # Common skill name mappings
        skill_map = {
            'python': 'Python',
            'java': 'Java',
            'javascript': 'JavaScript',
            'typescript': 'TypeScript',
            'c++': 'C++',
            'c#': 'C#',
            'node.js': 'Node.js',
            'nodejs': 'Node.js',
            'react': 'React',
            'angular': 'Angular',
            'vue': 'Vue',
            'django': 'Django',
            'flask': 'Flask',
            'fastapi': 'FastAPI',
            'sql': 'SQL',
            'mysql': 'MySQL',
            'postgresql': 'PostgreSQL',
            'mongodb': 'MongoDB',
            'redis': 'Redis',
            'aws': 'AWS',
            'azure': 'Azure',
            'gcp': 'GCP',
            'docker': 'Docker',
            'kubernetes': 'Kubernetes',
            'git': 'Git',
            'github': 'GitHub',
            'html': 'HTML',
            'css': 'CSS',
            'rest api': 'REST API',
            'graphql': 'GraphQL',
            'tensorflow': 'TensorFlow',
            'pytorch': 'PyTorch',
            'scikit-learn': 'Scikit-learn',
            'machine learning': 'Machine Learning',
            'deep learning': 'Deep Learning',
            'nlp': 'NLP',
            'computer vision': 'Computer Vision',
            'pandas': 'Pandas',
            'numpy': 'NumPy',
            'ci/cd': 'CI/CD',
            'devops': 'DevOps'
        }
        
        skill_lower = skill.lower()
        return skill_map.get(skill_lower, skill)
    
    def _extract_skills_section(self, text: str) -> List[str]:
        """
        Extract skills from dedicated skills section in resume
        
        Args:
            text: Resume text content
            
        Returns:
            List of skills from skills section
        """
        skills = []
        
        # Look for skills section
        skills_pattern = r'(?:skills?|technical skills?|core competencies|technologies)[\s:]*\n(.*?)(?:\n\n|\n[A-Z]|$)'
        matches = re.finditer(skills_pattern, text, re.IGNORECASE | re.DOTALL)
        
        for match in matches:
            skills_text = match.group(1)
            # Split by common delimiters
            skill_items = re.split(r'[,;•\n|]', skills_text)
            
            for item in skill_items:
                item = item.strip()
                # Remove bullet points and other markers
                item = re.sub(r'^[-•*]\s*', '', item)
                
                # Filter out non-skill items
                if self._is_valid_skill(item):
                    # Clean and normalize
                    item = self._normalize_skill_name(item)
                    skills.append(item)
        
        return skills
    
    def _is_valid_skill(self, text: str) -> bool:
        """
        Check if text is a valid skill (not a sentence or section header)
        
        Args:
            text: Text to validate
            
        Returns:
            True if valid skill, False otherwise
        """
        if not text or len(text) < 2 or len(text) > 50:
            return False
        
        # Exclude common non-skill patterns
        exclude_patterns = [
            r'^\d+\s+years?',  # "5 years"
            r'experience',  # Contains "experience"
            r'^\w+:$',  # Section headers like "SKILLS:"
            r':\s*$',  # Ends with colon (section header)
            r'at\s+\w+',  # "at Company"
            r'^\d{4}[-–]\d{4}',  # Date ranges
            r'bachelor|master|phd|degree',  # Education
            r'developed|implemented|worked|built',  # Action verbs
            r'senior|junior|lead|principal',  # Job titles
            r'engineer|developer|scientist|analyst',  # Job roles (unless part of skill)
            r'programming\s+languages',  # Section headers
            r'web\s+technologies',  # Section headers
            r'tools?\s+&?\s+technologies',  # Section headers
            r'databases?:',  # Section headers
            r'^projects?$',  # Section headers
            r'^technical\s+skills?$',  # Section headers
        ]
        
        text_lower = text.lower()
        for pattern in exclude_patterns:
            if re.search(pattern, text_lower):
                return False
        
        return True
    
    def extract_name(self, text: str) -> Optional[str]:
        """
        Extract candidate name from resume text
        
        Args:
            text: Resume text content
            
        Returns:
            Candidate name or None
        """
        lines = text.strip().split('\n')
        
        # Usually name is in first few lines
        for i, line in enumerate(lines[:5]):
            line = line.strip()
            # Skip empty lines and common headers
            if not line or line.upper() in ['RESUME', 'CV', 'CURRICULUM VITAE']:
                continue
            
            # Name is usually 2-4 words, capitalized, no special chars
            words = line.split()
            if 1 <= len(words) <= 4:
                # Check if it looks like a name (mostly alphabetic)
                if all(word.replace('.', '').isalpha() for word in words):
                    return line
        
        return None
    
    def extract_email(self, text: str) -> Optional[str]:
        """
        Extract email address from resume text
        
        Args:
            text: Resume text content
            
        Returns:
            Email address or None
        """
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, text)
        return match.group(0) if match else None
    
    def extract_phone(self, text: str) -> Optional[str]:
        """
        Extract phone number from resume text
        
        Args:
            text: Resume text content
            
        Returns:
            Phone number or None
        """
        # Indian phone number patterns
        phone_patterns = [
            r'\b\d{10}\b',  # 10 digits
            r'\b\d{5}\s?\d{5}\b',  # 5+5 digits
            r'\+91[-\s]?\d{10}\b',  # +91 prefix
            r'\b91\d{10}\b',  # 91 prefix
        ]
        
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        
        return None
    
    def extract_education(self, text: str) -> List[Dict[str, str]]:
        """
        Extract education details from resume text
        
        Args:
            text: Resume text content
            
        Returns:
            List of education entries
        """
        education = []
        seen_degrees = set()  # Track unique degrees
        
        # Look for education section
        edu_section_pattern = r'(?:ACADEMIC QUALIFICATION|EDUCATION|EDUCATIONAL QUALIFICATION)[\s:]*\n(.*?)(?:\n\n[A-Z]{2,}|\nWORK|EXPERIENCE|SKILLS|PROJECTS|$)'
        match = re.search(edu_section_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            edu_text = match.group(1)
            
            # Common degree patterns (more specific)
            degree_patterns = [
                (r'B\.?Tech\s+(?:CSE|CS|IT|ECE|EE|ME)?(?:\s+AI\s*&?\s*ML)?', 'B.Tech'),
                (r'Bachelor\s+of\s+Technology', 'B.Tech'),
                (r'M\.?Tech', 'M.Tech'),
                (r'Master\s+of\s+Technology', 'M.Tech'),
                (r'B\.?E\.?', 'B.E.'),
                (r'Bachelor\s+of\s+Engineering', 'B.E.'),
                (r'BCA', 'BCA'),
                (r'MCA', 'MCA'),
                (r'MBA', 'MBA'),
                (r'12th|XII|Higher\s+Secondary', '12th'),
                (r'10th|X(?:\s+|$)', '10th')
            ]
            
            for pattern, degree_name in degree_patterns:
                matches = re.finditer(pattern, edu_text, re.IGNORECASE)
                for match in matches:
                    degree_full = match.group(0).strip()
                    
                    # Skip if we've already seen this degree
                    if degree_name in seen_degrees:
                        continue
                    
                    seen_degrees.add(degree_name)
                    
                    # Get context around the match
                    context_start = max(0, match.start() - 50)
                    context_end = min(len(edu_text), match.end() + 150)
                    context = edu_text[context_start:context_end]
                    
                    # Extract year
                    year_match = re.search(r'(20\d{2}|19\d{2})', context)
                    year = year_match.group(0) if year_match else None
                    
                    # Extract percentage/CGPA
                    score_match = re.search(r'(\d+\.?\d*)\s*(?:%|CGPA|cgpa)', context, re.IGNORECASE)
                    score = score_match.group(0) if score_match else None
                    
                    # Extract institution (look for University, College, Board)
                    institution = None
                    inst_match = re.search(r'([A-Z][A-Za-z\s]+(?:University|College|Board|Institute))', context)
                    if inst_match:
                        institution = inst_match.group(1).strip()
                    
                    education.append({
                        'degree': degree_full,
                        'institution': institution,
                        'year': year,
                        'score': score
                    })
        
        return education
    
    def extract_work_experience(self, text: str) -> List[Dict[str, str]]:
        """
        Extract work experience from resume text
        
        Args:
            text: Resume text content
            
        Returns:
            List of work experience entries
        """
        experiences = []
        seen_entries = set()  # Track unique entries
        
        # Look for experience section
        exp_section_pattern = r'(?:WORK EXPERIENCE|EXPERIENCE|PROFESSIONAL EXPERIENCE|EMPLOYMENT)[\s:]*\n(.*?)(?:\n\n[A-Z]{2,}|\nEDUCATION|SKILLS|PROJECTS|$)'
        match = re.search(exp_section_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            exp_text = match.group(1)
            
            # Split by lines and look for job entries
            lines = exp_text.split('\n')
            current_entry = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Look for "Title at Company (Year-Year)" pattern
                exp_match = re.search(r'((?:Software|Senior|Junior|Lead|Full Stack|Backend|Frontend|Data|ML|DevOps)?\s*(?:Engineer|Developer|Analyst|Manager|Intern))\s+at\s+([A-Za-z\s&,\.]+?)(?:\s*\()(\d{4}\s*[-–]\s*(?:\d{4}|Present|Current))', line, re.IGNORECASE)
                
                if exp_match:
                    title = exp_match.group(1).strip()
                    company = exp_match.group(2).strip()
                    duration = exp_match.group(3).strip()
                    
                    # Create unique key
                    entry_key = f"{title}|{company}"
                    
                    if entry_key not in seen_entries:
                        seen_entries.add(entry_key)
                        experiences.append({
                            'title': title,
                            'company': company,
                            'duration': duration
                        })
        
        return experiences
    
    def extract_experience_years(self, text: str) -> Optional[float]:
        """
        Extract years of experience from resume text
        
        Args:
            text: Resume text content
            
        Returns:
            Years of experience or None
        """
        # Look for patterns like "5 years of experience", "5+ years", etc.
        patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience[:\s]+(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s+(?:of\s+)?experience'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return float(match.group(1))
                except ValueError:
                    continue
        
        return None
    
    def parse_resume_full(self, file_content: bytes, filename: str) -> Dict:
        """
        Parse resume and extract all information
        
        Args:
            file_content: File content as bytes
            filename: Name of the file
            
        Returns:
            Dictionary with extracted information
        """
        # Extract text
        text = self.parse_resume(file_content, filename)
        
        # Extract all information
        name = self.extract_name(text)
        email = self.extract_email(text)
        phone = self.extract_phone(text)
        skills = self.extract_skills_from_text(text)
        education = self.extract_education(text)
        work_experience = self.extract_work_experience(text)
        experience_years = self.extract_experience_years(text)
        
        return {
            "text": text,
            "name": name,
            "email": email,
            "phone": phone,
            "skills": skills,
            "education": education,
            "work_experience": work_experience,
            "experience_years": experience_years,
            "filename": filename
        }
